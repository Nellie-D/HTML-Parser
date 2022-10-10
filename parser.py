#This python project is a learning experience 
#Please note that although I could have used
#pip - python-docx, it would've changed the entire
#way this was coded. With this code, imports can be made to
#alternative types of text documents such as .txt and .pdf

from encodings import utf_8
from html.parser import HTMLParser
import unicodedata
import codecs
import os
from docx import Document

coding: utf_8

newfilename = input("Enter the desired name of your document: ")
concat = input("What file type do you wish to create?")
#note that the pc parser is not necessary with the newly improved ios version

class parseHTMLDataPC(HTMLParser):
    
    def handle_starttag(self, tag, attrs):
      #  print("Encountered a tag:", tag)
        
        if (tag) == 'p':

            (line1, column) = self.getpos()
    
    def handle_data(self, data):
        newFilePath = os.path.abspath(newfilename)
        
        newList = []
        if concat == ".docx":
            doc = Document()
          
            (line, column) = self.getpos()
            if line > 668:
                doc.add_heading(newfilename)
                
                for n in data:
                    newList.append(n)
                i = 0
                listLength = len(newList)
                while i < listLength:
                    doc.add_paragraph(newList[i].encode('ascii', 'ignore').decode())
                    
                    i += 1
                
                
             
            #newSave = open(newFilePath + concat, "a")
            doc.save(newFilePath + concat)
            
        else:
            newSave = open(newFilePath + ".doc", "a")
            (line, column) = self.getpos()
            #notion documents end css formatting before line 668
            #greater than line 668, the document data is found
            
            if line > 668:

                #adds a tab to the beginning of each newline
                newSave.writelines("\t" + data + "\n")

            newSave.flush()  
            newSave.close()



class parseHTMLDataiOS(HTMLParser):
    
    def handle_starttag(self, tag, attrs):
        print("Encountered a tag:", tag)
        if (tag) == 'p':
            (line1, column) = self.getpos()
           
    def handle_data(self, data):

        newFilePath = os.path.abspath(newfilename)

        #newSave = open(newFilePath + ".docx", "a")
        newSave = open(newFilePath + concat, "a", encoding="utf-8")
        (line, column) = self.getpos()
        newlineType = '\n'
        tabType = "\t"
        myDataList = []
        
        if line > 668:
            
             
            for n in data:
            
                myDataList.append(n)
                if n == newlineType:
                    myDataList.append("\t")
              
            
            if newlineType in myDataList and tabType not in myDataList:
                #returns true
                newLineIndex = myDataList.index(newlineType)
                print(newLineIndex)
                paragraphStartIndex = newLineIndex + 1
                myDataList.insert(paragraphStartIndex, "\t")
              
           
            else:
                #this helps to break up blocks of text that are formatted differenlty based 
                #on how the document was written (ex. if document was written both pc and ios)
                #myDataList.insert(0, "\n")
                
                if tabType not in myDataList:

                    myDataList.insert(0, "\t")
                    myDataList.append("\n")
                else: 
                    myDataList.append("\n")
                print(False)

            print(myDataList)
            newSave.writelines(myDataList)
                    
                        
            
        newSave.flush()
        newSave.close()

def main(): 

    importType = float(input("Was your document written on PC(1) or on iOS(2) device?"))
    downloadedFile = input("Enter the name of the exported file: ")
    
    myStore = downloadedFile
    downloadPath = os.path.abspath(myStore)
    f = codecs.open(downloadPath, encoding="utf8")


    if importType == 1:
        parserprint = parseHTMLDataPC()
        parserprint.feed(f.read())
    elif importType == 2:
        parserprint = parseHTMLDataiOS()
        parserprint.feed(f.read())
    else:
        print("Invalid entry: " + importType)
        

    
    parserprint.close()
    f.close()
    

   


main()